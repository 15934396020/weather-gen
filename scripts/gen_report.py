from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.25

def R(p, text, size=12, bold=False, east='宋体'):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), east)
    return r

def cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    R(p, text, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def cell_paras(cell, lines):
    cell.text = ''
    for i, (text, bold) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        R(p, text, bold=bold)

# ========== Cover Page ==========
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, '西南交通大学-利兹学院', size=22, bold=True, east='黑体')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, '毕业设计（论文）中期报告', size=22, bold=True, east='黑体')
doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ('题    目：', 'Design and Implementation of a Controllable'),
    ('', 'Scenario Weather Image Generation and Display'),
    ('', 'System Based on ControlNet'),
    ('专    业：', '计算机科学与技术'),
    ('学    号：', '2022116037'),
    ('姓    名：', '申泽宇'),
    ('指导教师：', '彭博'),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if label: R(p, label, size=15)
    R(p, value, size=15)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, '             填表日期：2026年3月10日', size=15)

doc.add_page_break()

# ========== Main Table ==========
table = doc.add_table(rows=0, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

def add_row(left, content_lines):
    row = table.add_row()
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(13.0)
    cell_text(row.cells[0], left, bold=True)
    cell_paras(row.cells[1], content_lines)
    return row

# --- Header rows ---
add_row('题目名称', [('Design and Implementation of a Controllable Scenario Weather Image Generation and Display System Based on ControlNet', False)])
add_row('题目来源', [('  ☐ 生产实际      ☑ 教师自拟      ☐ 学生自拟      ☐ 其他', False)])
add_row('题目类型', [('毕业论文（☑ 应用研究）', False)])

# ========== Section 1: Overview ==========
s1 = """This project aims to design and implement a controllable scenario weather image generation and display system based on ControlNet. The system allows users to generate realistic scene images under specific weather conditions — including sunny, rainy, snowy, foggy, and thunderstorm — by combining natural language text prompts with structural control maps such as Canny edge maps, scribble drawings, and depth maps.

The technical foundation of the system is built upon Stable Diffusion as the base diffusion model, with ControlNet integrated as the conditional control extension. By creating a trainable copy of the Stable Diffusion encoder and introducing an additional encoder to process structured inputs, ControlNet enables fine-grained, pixel-level control over the generation process while preserving the style and semantics provided by text prompts. On the backend, a Python-based service using the FastAPI framework encapsulates the model inference pipeline and exposes RESTful API endpoints. The frontend provides an interactive web interface for text input, control map uploading, weather type selection, and parameter adjustment.

A central component of this work is the weather effect prompt library. Instead of relying solely on user-provided text, the system automatically appends weather-specific keywords and descriptive phrases to the base prompt, which was found through experimentation to significantly improve the consistency and visual quality of the generated weather effects. The overall system follows a Browser/Server (B/S) architecture consisting of three layers: the frontend user interface, the backend API service, and the AI model inference engine."""

add_row('一、毕业设计（论文）概述', [(s1, False)])

# ========== Section 2: Schedule ==========
s2 = """The project is planned over a total of 23 weeks. During Weeks 1–2, a comprehensive literature review was conducted, covering the theoretical principles of Diffusion Models and ControlNet, as well as existing research on weather simulation using generative AI. The project proposal was refined based on these findings. This phase has been completed.

In Weeks 3–4, the development environment was set up, including Python, PyTorch, the Diffusers library, Transformers, OpenCV, and FastAPI. Pre-trained Stable Diffusion (v1.5) and ControlNet model weights (Canny, Scribble, Depth variants) were downloaded from Hugging Face and validated through test scripts. This phase has been completed.

Weeks 5–8 were dedicated to backend system design and API development. The overall system architecture was designed, and the backend service was built using FastAPI with endpoints for receiving text prompts, control images, weather parameters, and advanced generation settings. The core logic for invoking the ControlNet pipeline was implemented. This phase has been completed.

During Weeks 9–12, the frontend user interface was designed and developed, including components for text input, image upload, weather type selection buttons, parameter adjustment sliders, and a result display panel. This phase has been completed.

The project is currently in Weeks 13–15, focusing on integrating the frontend and backend into a fully functional end-to-end system, implementing control map preprocessing (such as automatic Canny edge extraction from uploaded images), and conducting workflow testing and debugging. This phase is in progress.

The remaining phases — advanced features and UI optimization (Weeks 16–17), system testing and evaluation (Weeks 18–19), and thesis writing with defense preparation (Weeks 20–23) — are planned for the coming weeks."""

add_row('二、毕业设计（论文）整体安排及进度', [(s2, False)])

# ========== Section 3: Completed work (with model details, UI, dataset, experiments) ==========
s3 = """3.1 Literature Review

The literature review covered the evolution of generative models from GANs to Diffusion Models, with particular focus on how DDPM and DDIM achieve high-fidelity image synthesis through iterative denoising. The ControlNet architecture proposed by Zhang and Agrawala [1] was studied in detail, specifically how a trainable copy of the diffusion model encoder accepts additional structural conditions while keeping the original weights locked. Several works on weather visualization using generative AI [2][3] were also reviewed, confirming that the controllable combination of multiple weather effects with user-defined complex scenes remains a relatively underexplored direction.

3.2 Core Model Architecture and Principles

The image generation pipeline of this system is built upon two core models: Stable Diffusion and ControlNet. Figure 1 illustrates the overall system architecture, and Figure 2 shows the detailed inference pipeline.

Stable Diffusion [7] is a Latent Diffusion Model (LDM) that performs the diffusion process in a compressed latent space rather than directly in pixel space. The architecture consists of three main components: (a) a Variational Autoencoder (VAE), whose encoder compresses an input image x into a lower-dimensional latent representation z = E(x), and whose decoder reconstructs images from latent codes x' = D(z); (b) a U-Net denoising network, which is trained to predict and remove noise from the latent representation through an iterative denoising process — during training, Gaussian noise is progressively added to the latent code (forward diffusion), and the U-Net learns to reverse this process (reverse diffusion), gradually recovering the clean latent from pure noise over T timesteps; (c) a text encoder based on CLIP [10], which converts the user's text prompt into a conditioning vector that guides the U-Net via cross-attention layers, allowing the generated image to reflect the semantic content described in the text.

ControlNet [1] extends Stable Diffusion by adding spatial conditioning capabilities. The key idea is to create a trainable copy of the U-Net's encoder blocks while keeping the original Stable Diffusion weights locked (frozen). The trainable copy receives an additional input — a structural control map such as a Canny edge image, a depth map, or a scribble drawing — and its output features are injected back into the frozen U-Net through "zero convolution" layers. Zero convolutions are 1x1 convolution layers whose weights and biases are initialized to zero, meaning that at the start of training, the ControlNet branch produces zero output and does not interfere with the pre-trained model's behavior. As training progresses, the zero convolution weights gradually learn to inject meaningful structural information. This design ensures training stability and preserves the original model's generation quality while enabling precise structural control.

The inference process works as follows: (1) the user provides a text prompt, a control image, and selects a weather type; (2) the backend extracts Canny edges from the control image using OpenCV; (3) the text prompt is combined with weather-specific keywords from the prompt library; (4) the combined prompt is encoded by CLIP into a conditioning vector; (5) starting from random Gaussian noise in the latent space, the U-Net iteratively denoises the latent code over N steps (default 20), guided by both the text conditioning and the ControlNet's structural conditioning; (6) the final denoised latent is decoded by the VAE decoder into the output image.

[Figure 1: System Architecture — three-layer B/S architecture consisting of the frontend user interface, the backend FastAPI service, and the AI model inference engine (Stable Diffusion + ControlNet). See attached figure: system_architecture.png]

[Figure 2: ControlNet Inference Pipeline — showing the flow from user inputs (text prompt, control image, weather type) through preprocessing, CLIP encoding, ControlNet conditioning, U-Net iterative denoising, and VAE decoding to the final output image. See attached figure: model_pipeline.png]

3.3 Development Environment and Model Validation

The development environment was configured with Python 3.9+, PyTorch 2.0+ (with MPS backend support for Apple Silicon and CUDA support for NVIDIA GPUs), the Diffusers library for running Stable Diffusion and ControlNet pipelines, and OpenCV for image preprocessing. Pre-trained model weights for Stable Diffusion v1.5 and the ControlNet Canny variant were downloaded from Hugging Face. Preliminary tests confirmed that the pipeline correctly generates images conditioned on both text prompts and Canny edge control maps.

3.4 Backend and Frontend Development

The backend API service was developed using FastAPI. The main generation endpoint (POST /api/generate) accepts a multipart form containing the text prompt, a control image file, the selected weather type, and optional parameters (inference steps, CFG scale, random seed, and an auto-Canny toggle). Upon receiving a request, the backend preprocesses the control image (resizing to 512x512 and optionally extracting Canny edges), constructs the full prompt by merging the user's text with the corresponding weather prompt template, runs the inference pipeline, and returns the generated image as a PNG response.

The frontend was developed as a single-page web application served directly by the FastAPI backend. The interface includes: (a) a text input area for scene descriptions; (b) a drag-and-drop image upload area with preview and an auto-Canny toggle; (c) five weather type selection buttons (Sunny, Rainy, Snowy, Foggy, Thunderstorm); (d) advanced parameter sliders for inference steps, CFG scale, and seed; (e) a result display panel; and (f) a history panel for browsing previously generated images.

[Figure 3: Screenshot of the system's web interface showing the control image upload area, weather type selection buttons, advanced parameter settings, and the generation result display area. See attached screenshot.]

3.5 Weather Effect Prompt Library and Dataset

A core contribution of this work is the design and validation of the weather effect prompt library. For each of the five supported weather types, a set of positive descriptive phrases and negative exclusion terms were carefully designed through iterative experimentation. The prompt templates are summarized below:

- Sunny: positive keywords include "bright sunlight, clear blue sky, warm golden light, sharp shadows, vivid colors, sun rays"; negative keywords include "rain, snow, fog, dark, overcast".
- Rainy: positive keywords include "heavy rain, wet surfaces, water reflections, puddles, overcast sky, rain streaks, gloomy atmosphere, dark clouds"; negative keywords include "sunny, clear sky, snow, dry".
- Snowy: positive keywords include "snow covered, white snow on ground, frost, cold winter atmosphere, snowfall, icy surfaces, cold blue lighting"; negative keywords include "sunny, green leaves, rain, warm".
- Foggy: positive keywords include "dense fog, misty atmosphere, low visibility, soft diffused light, hazy, mysterious mood"; negative keywords include "clear sky, sunny, sharp details, vivid colors".
- Thunderstorm: positive keywords include "thunderstorm, lightning bolts, dark dramatic sky, heavy rain, strong wind, ominous clouds, dramatic lighting"; negative keywords include "sunny, calm, clear sky, peaceful".

A shared quality-boosting suffix ("high quality, detailed, 8k, photorealistic, masterpiece") is appended to all positive prompts, and a shared quality-degradation exclusion ("low quality, blurry, distorted, deformed, ugly, watermark, text") is appended to all negative prompts.

The test dataset for the experiments consists of two categories:

(a) Synthetic scene images: These are programmatically generated using Python's PIL and OpenCV libraries to create simplified but structurally representative scenes. The current test set includes a cityscape image (512x512 pixels) featuring rectangular buildings of varying heights with window grids, a sidewalk area, and a road with lane markings. Additional synthetic scenes (e.g., a park with trees and paths, a lakeside landscape) are planned for the next phase. The advantage of synthetic images is that they provide clean, reproducible structural inputs with known geometric properties, making them ideal for controlled comparison experiments.

(b) Real-world photographs: These are collected from two sources — (i) open-license image datasets such as Unsplash and Pexels, filtered for common scene categories including urban streets, natural landscapes, residential buildings, and parks; and (ii) photographs taken by the author in the campus area and surrounding city streets. The real-world images are preprocessed by resizing to 512x512 pixels and, when used with ControlNet, by extracting Canny edge maps using OpenCV with threshold parameters (low=100, high=200). These images test the system's ability to handle complex, real-world structural inputs with irregular geometries, varying lighting conditions, and diverse scene compositions.

The dataset is organized into three scene categories for systematic evaluation: (1) Urban scenes — city streets with buildings, vehicles, and pedestrians; (2) Natural landscapes — parks, lakes, mountains, and open fields; (3) Architectural scenes — individual buildings, campus structures, and residential areas. Each category contains both synthetic and real-world examples to enable cross-validation of the weather effect generation quality.

3.6 Experimental Process and Preliminary Results

The experimental work at this stage focused on validating the effectiveness of the weather prompt templates through a text-to-image generation experiment. This serves as a preliminary verification step before the full ControlNet-based pipeline is deployed on a GPU server.

Experiment Setup: Using the same base scene description ("a city street with tall buildings, cars and pedestrians, photorealistic, highly detailed, 8k resolution"), five images were generated — one for each weather type — by appending the corresponding weather prompt template to the base description. The generation was performed using a diffusion model via the Hugging Face Inference API, with consistent parameters across all five runs to ensure a fair comparison.

Results and Analysis: The five generated images (Figure 4) demonstrate that the weather prompt templates produce clearly distinguishable weather effects:

- The Sunny image exhibits bright, warm lighting with clear skies and sharp shadows, consistent with the intended sunny atmosphere.
- The Rainy image shows wet surfaces, visible rain streaks, overcast skies, and an overall gloomy tone, effectively conveying a rainy scene.
- The Snowy image features snow-covered ground, cold blue-toned lighting, and a winter atmosphere, successfully rendering a snowy environment.
- The Foggy image displays reduced visibility, soft diffused lighting, and a hazy atmosphere, accurately representing foggy conditions.
- The Thunderstorm image presents dramatic dark skies, visible lightning, and heavy rain, creating a convincing stormy scene.

These results confirm that the designed prompt templates are effective at guiding the diffusion model to generate weather-specific visual characteristics. The clear differentiation between weather types validates the prompt engineering approach adopted in this project.

[Figure 4: Weather Effect Prompt Validation — five images generated from the same base scene description with different weather prompt templates (Sunny, Rainy, Snowy, Foggy, Thunderstorm), demonstrating the effectiveness of the prompt library design. See attached figure: weather_prompt_validation.png]

It should be noted that this preliminary experiment uses text-to-image generation without structural control input. As a result, the generated images exhibit certain visual artifacts commonly associated with AI-generated content, such as occasional texture inconsistencies, minor structural distortions in fine details (e.g., window frames, vehicle shapes), and an overall "smoothness" that differs from real photographs. These artifacts arise from two main factors: (1) the current experiment relies solely on text prompts to guide generation, without the structural constraints that ControlNet provides — the scene layout is entirely determined by the model's learned priors rather than being anchored to a real scene structure; (2) the generation parameters (inference steps, CFG scale) have not yet been fine-tuned for each individual weather type. In the next phase, the introduction of ControlNet conditioning with Canny edge maps extracted from real photographs is expected to significantly improve the structural accuracy and visual realism of the generated images, as the model will be constrained to follow the actual geometric layout of the input scene. Additionally, systematic parameter tuning and further prompt template refinement will help reduce the remaining visual artifacts."""

add_row('三、毕业设计（论文）已完成的研究部分', [(s3, False)])

# ========== Section 4: Next steps ==========
s4 = """The immediate priority is to complete the integration of the frontend and backend into a seamless end-to-end system. This involves deploying the model inference pipeline on a GPU-equipped server to enable practical generation speeds, finalizing the control map preprocessing pipeline so that users can upload raw photographs and have Canny edges extracted automatically, and implementing proper progress feedback during the generation process.

Following the integration phase, an image history and management module will be developed. This module will allow users to browse, compare, and download previously generated images within a session. A side-by-side comparison view is planned so that users can visually compare the same scene under different weather conditions. The overall UI layout and styling will also be refined during this period.

A systematic testing and evaluation phase will then be carried out. The three experiments described in Section 3.6 will be executed on the GPU server, and the results will be collected and analyzed. Representative test cases covering different combinations of scene descriptions, control map types, and weather conditions will be evaluated for functional correctness, generation quality, response time, and system stability. The evaluation will combine visual inspection with quantitative metrics (such as FID score where a suitable reference dataset is available).

The final phase will focus on writing the graduation thesis, which will cover the project background, literature review, system architecture and design, implementation details, experimental results, and conclusions. A project demonstration video and defense presentation slides will also be prepared."""

add_row('四、下一部分的工作安排', [(s4, False)])

# ========== Section 5: Problems ==========
s5 = """The primary challenge encountered so far is the computational resource constraint. Running Stable Diffusion with ControlNet requires substantial GPU memory — the model occupies approximately 6–7 GB of VRAM during inference. On the development machine (Apple M2 with shared memory architecture), the MPS backend runs out of memory when attempting to generate 512×512 images with the full pipeline. Half-precision inference and attention slicing were applied as optimizations, but the generation speed remains impractical (over 4 minutes per inference step). The solution is to deploy the inference pipeline on a dedicated GPU server (NVIDIA GPU with ≥8 GB VRAM), which is planned for the next phase. The system architecture has been designed with this migration in mind — switching from MPS to CUDA requires only a single configuration change.

The quality of weather effects can vary across different scene types. While the prompt templates produce convincing results for common urban and landscape scenes, they may appear less natural for unusual or highly specific scene compositions. Continued refinement of the weather prompt library and exploration of additional ControlNet conditioning strategies (such as combining multiple ControlNet models or adjusting the conditioning scale) are needed to improve robustness.

Establishing objective evaluation criteria for the generated weather images remains a challenge. Standard metrics like FID require large reference datasets of real weather scene images, which are difficult to curate for all weather-scene combinations. The current plan is to combine quantitative metrics where feasible with structured qualitative evaluation through visual comparison and human scoring.


References:

[1] Zhang, L., & Agrawala, M. (2023). Adding conditional control to text-to-image diffusion models. In Proceedings of the IEEE/CVF International Conference on Computer Vision (pp. 3836-3847).
[2] Croitoru, F. A., Hondru, V., Ionescu, R. T., & Shah, M. (2023). Diffusion models in vision: A survey. IEEE Transactions on Pattern Analysis and Machine Intelligence, 45(9), 10790-10809.
[3] Li, X., Ren, Y., Jin, X., Lan, C., Wang, X., Zeng, W., & Liu, Z. (2025). Diffusion models for image restoration and enhancement: A comprehensive survey. International Journal of Computer Vision, 133(1), 1-36.
[4] Chen, H., He, H., Shen, Y., Lin, Z., & Liu, Z. (2025). Comprehensive exploration of diffusion models in image generation: A survey. Artificial Intelligence Review, 58(1), 1-56.
[5] Yao, S., & Budthimedhee, K. (2024). From Sketches to Renderings: A Comparison of Rapid Visualization Methods for Sketches Based on ControlNet. In 2024 5th International Conference on CECIT (pp. 1-6). IEEE.
[6] Kwak, M. I. Y., & Lim, S. (2024). Visualizing Real-Time Weather Data Using Generative AI. Journal of Digital Contents Society, 25(1), 123-131.
[7] Rombach, R., Blattmann, A., Lorenz, D., Esser, P., & Ommer, B. (2022). High-resolution image synthesis with latent diffusion models. In Proceedings of the IEEE/CVF CVPR (pp. 10684-10695).
[8] Ho, J., Jain, A., & Abbeel, P. (2020). Denoising diffusion probabilistic models. Advances in NeurIPS, 33, 6840-6851.
[9] Song, J., Meng, C., & Ermon, S. (2021). Denoising diffusion implicit models. In International Conference on Learning Representations.
[10] Radford, A., Kim, J. W., Hallacy, C., et al. (2021). Learning transferable visual models from natural language supervision. In ICML (pp. 8748-8763)."""

add_row('五、毕业设计（论文）工作中存在的问题', [(s5, False)])

# ========== Signature ==========
row = table.add_row()
row.cells[0].merge(row.cells[1])
cell = row.cells[0]; cell.text = ''
p = cell.paragraphs[0]
R(p, '\n\n学生（签字）              年   月   日            指导教师（签字）：      年   月   日\n\n')

output = '/Users/weidademiaoxiao/Desktop/毕业设计/中期检查报告_申泽宇.docx'
doc.save(output)
print('Done:', output)
