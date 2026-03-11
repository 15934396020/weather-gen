"""Generate final report with all images embedded"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import os

BASE = os.path.expanduser("~/Desktop/毕业设计")
FIGDIR = os.path.join(BASE, "weather-gen/docs/figures")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.25

def R(p, text, size=12, bold=False, east='宋体'):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), east)
    return r

def add_figure(doc, img_path, caption, width=Cm(15)):
    """Add a figure with caption to the document"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(cap, caption, size=10, bold=False)
    doc.add_paragraph()  # spacing

def cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    R(p, text, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def cell_paras(cell, lines):
    cell.text = ''
    for i, (text, bold) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        R(p, text, bold=bold)

# ========== Cover Page ==========
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, '西南交通大学-利兹学院', size=22, bold=True, east='黑体')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, '毕业设计（论文）中期报告', size=22, bold=True, east='黑体')
doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ('题    目：', 'Design and Implementation of a Controllable'),
    ('', 'Scenario Weather Image Generation and Display'),
    ('', 'System Based on ControlNet'),
    ('专    业：', '计算机科学与技术'),
    ('学    号：', '2022116037'),
    ('姓    名：', '申泽宇'),
    ('指导教师：', '彭博'),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if label: R(p, label, size=15)
    R(p, value, size=15)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, '             填表日期：2026年3月10日', size=15)

doc.add_page_break()

# ========== Main Table ==========
table = doc.add_table(rows=0, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

def add_row(left, content_lines):
    row = table.add_row()
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(13.0)
    cell_text(row.cells[0], left, bold=True)
    cell_paras(row.cells[1], content_lines)
    return row

def add_row_with_images(left, text_before, images, text_after):
    """Add a row that contains text, then images, then more text"""
    row = table.add_row()
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(13.0)
    cell_text(row.cells[0], left, bold=True)

    cell = row.cells[1]
    cell.text = ''

    # Text before images
    for i, (text, bold) in enumerate(text_before):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        R(p, text, bold=bold)

    # Images
    for img_path, caption in images:
        if os.path.exists(img_path):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Cm(12))
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            R(cap, caption, size=10)
            cell.add_paragraph()  # spacing

    # Text after images
    for text, bold in text_after:
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        R(p, text, bold=bold)

    return row

# --- Header rows ---
add_row('题目名称', [('Design and Implementation of a Controllable Scenario Weather Image Generation and Display System Based on ControlNet', False)])
add_row('题目来源', [('  ☐ 生产实际      ☑ 教师自拟      ☐ 学生自拟      ☐ 其他', False)])
add_row('题目类型', [('毕业论文（☑ 应用研究）', False)])

# ========== Section 1 ==========
s1 = """This project aims to design and implement a controllable scenario weather image generation and display system based on ControlNet. The system allows users to generate realistic scene images under specific weather conditions — including sunny, rainy, snowy, foggy, and thunderstorm — by combining natural language text prompts with structural control maps such as Canny edge maps, scribble drawings, and depth maps.

The technical foundation of the system is built upon Stable Diffusion as the base diffusion model, with ControlNet integrated as the conditional control extension. By creating a trainable copy of the Stable Diffusion encoder and introducing an additional encoder to process structured inputs, ControlNet enables fine-grained, pixel-level control over the generation process while preserving the style and semantics provided by text prompts. On the backend, a Python-based service using the FastAPI framework encapsulates the model inference pipeline and exposes RESTful API endpoints. The frontend provides an interactive web interface for text input, control map uploading, weather type selection, and parameter adjustment.

A central component of this work is the weather effect prompt library. Instead of relying solely on user-provided text, the system automatically appends weather-specific keywords and descriptive phrases to the base prompt, which was found through experimentation to significantly improve the consistency and visual quality of the generated weather effects. The overall system follows a Browser/Server (B/S) architecture consisting of three layers: the frontend user interface, the backend API service, and the AI model inference engine."""

add_row('一、毕业设计（论文）概述', [(s1, False)])

# ========== Section 2 ==========
s2 = """The project is planned over a total of 23 weeks. During Weeks 1–2, a comprehensive literature review was conducted, covering the theoretical principles of Diffusion Models and ControlNet, as well as existing research on weather simulation using generative AI. The project proposal was refined based on these findings. This phase has been completed.

In Weeks 3–4, the development environment was set up, including Python, PyTorch, the Diffusers library, Transformers, OpenCV, and FastAPI. Pre-trained Stable Diffusion (v1.5) and ControlNet model weights (Canny, Scribble, Depth variants) were downloaded from Hugging Face and validated through test scripts. This phase has been completed.

Weeks 5–8 were dedicated to backend system design and API development. The overall system architecture was designed, and the backend service was built using FastAPI with endpoints for receiving text prompts, control images, weather parameters, and advanced generation settings. The core logic for invoking the ControlNet pipeline was implemented. This phase has been completed.

During Weeks 9–12, the frontend user interface was designed and developed, including components for text input, image upload, weather type selection buttons, parameter adjustment sliders, and a result display panel. This phase has been completed.

The project is currently in Weeks 13–15, focusing on integrating the frontend and backend into a fully functional end-to-end system, implementing control map preprocessing (such as automatic Canny edge extraction from uploaded images), and conducting workflow testing and debugging. This phase is in progress.

The remaining phases — advanced features and UI optimization (Weeks 16–17), system testing and evaluation (Weeks 18–19), and thesis writing with defense preparation (Weeks 20–23) — are planned for the coming weeks."""

add_row('二、毕业设计（论文）整体安排及进度', [(s2, False)])

# ========== Section 3 (with images) ==========
s3_before = [
    ("""3.1 Literature Review

The literature review covered the evolution of generative models from GANs to Diffusion Models, with particular focus on how DDPM and DDIM achieve high-fidelity image synthesis through iterative denoising. The ControlNet architecture proposed by Zhang and Agrawala [1] was studied in detail, specifically how a trainable copy of the diffusion model encoder accepts additional structural conditions while keeping the original weights locked. Several works on weather visualization using generative AI [2][3] were also reviewed, confirming that the controllable combination of multiple weather effects with user-defined complex scenes remains a relatively underexplored direction.

3.2 Core Model Architecture and Principles

The image generation pipeline of this system is built upon two core models: Stable Diffusion and ControlNet. Figure 1 illustrates the overall system architecture, and Figure 2 shows the detailed inference pipeline.

Stable Diffusion [7] is a Latent Diffusion Model (LDM) that performs the diffusion process in a compressed latent space rather than directly in pixel space. The architecture consists of three main components: (a) a Variational Autoencoder (VAE), whose encoder compresses an input image x into a lower-dimensional latent representation z = E(x), and whose decoder reconstructs images from latent codes; (b) a U-Net denoising network, which is trained to predict and remove noise from the latent representation through an iterative denoising process; (c) a text encoder based on CLIP [10], which converts the user's text prompt into a conditioning vector that guides the U-Net via cross-attention layers.

ControlNet [1] extends Stable Diffusion by adding spatial conditioning capabilities. The key idea is to create a trainable copy of the U-Net's encoder blocks while keeping the original Stable Diffusion weights locked (frozen). The trainable copy receives an additional input — a structural control map such as a Canny edge image — and its output features are injected back into the frozen U-Net through "zero convolution" layers initialized to zero. This design ensures training stability and preserves the original model's generation quality while enabling precise structural control.""", False),
]

images_1 = [
    (os.path.join(FIGDIR, "system_architecture.png"), "Figure 1: System Architecture — three-layer B/S architecture (Frontend, Backend API, AI Model Engine)"),
    (os.path.join(FIGDIR, "model_pipeline.png"), "Figure 2: ControlNet Inference Pipeline — from user inputs through preprocessing, encoding, denoising, to output"),
]

s3_mid = [
    ("""3.3 Development Environment and Model Validation

The development environment was configured with Python 3.9+, PyTorch 2.0+ (with MPS backend support for Apple Silicon and CUDA support for NVIDIA GPUs), the Diffusers library for running Stable Diffusion and ControlNet pipelines, and OpenCV for image preprocessing. Pre-trained model weights for Stable Diffusion v1.5 and the ControlNet Canny variant were downloaded from Hugging Face and validated.

3.4 Backend and Frontend Development

The backend API service was developed using FastAPI. The main generation endpoint (POST /api/generate) accepts a multipart form containing the text prompt, a control image file, the selected weather type, and optional parameters (inference steps, CFG scale, random seed, and an auto-Canny toggle). The frontend was developed as a single-page web application with a left-right split layout: the left panel contains input controls (scene description, control image upload, weather type selection, advanced parameters), and the right panel displays the generated result and generation history. Figure 3 shows the system interface.""", False),
]

images_2 = [
    (os.path.join(BASE, "中期报告页面展示图.png"), "Figure 3: System Web Interface — left panel for input controls, right panel for generated result display"),
]

s3_dataset = [
    ("""3.5 Weather Effect Prompt Library and Dataset

A core contribution of this work is the design and validation of the weather effect prompt library. For each of the five supported weather types, a set of positive descriptive phrases and negative exclusion terms were carefully designed through iterative experimentation:

- Sunny: "bright sunlight, clear blue sky, warm golden light, sharp shadows, vivid colors, sun rays"
- Rainy: "heavy rain, wet surfaces, water reflections, puddles, overcast sky, rain streaks, gloomy atmosphere"
- Snowy: "snow covered, white snow on ground, frost, cold winter atmosphere, snowfall, cold blue lighting"
- Foggy: "dense fog, misty atmosphere, low visibility, soft diffused light, hazy, mysterious mood"
- Thunderstorm: "thunderstorm, lightning bolts, dark dramatic sky, heavy rain, ominous clouds, dramatic lighting"

The test dataset consists of two categories: (a) Synthetic scene images generated using Python's PIL and OpenCV libraries (512x512 pixels), including a cityscape with buildings, windows, roads, and sidewalks, providing clean and reproducible structural inputs; (b) Real-world photographs collected from open-license image datasets (Unsplash, Pexels) and the author's own photography of campus and city scenes, preprocessed by resizing to 512x512 and extracting Canny edge maps (threshold low=100, high=200). The dataset is organized into three scene categories: urban scenes, natural landscapes, and architectural scenes.

3.6 Experimental Process and Preliminary Results

The experimental work at this stage focused on validating the effectiveness of the weather prompt templates through a text-to-image generation experiment, serving as a preliminary verification step before the full ControlNet pipeline is deployed on a GPU server.

Experiment Setup: Using the same base scene description ("a city street with tall buildings, cars and pedestrians"), five images were generated — one for each weather type — by appending the corresponding weather prompt template. The generation was performed using a diffusion model via the Hugging Face Inference API, with consistent parameters across all five runs.""", False),
]

images_3 = [
    (os.path.join(FIGDIR, "experiment_results_grid.png"), "Figure 4: Weather Effect Generation Results — five images generated from the same base scene description under different weather conditions, with generation parameters shown"),
]

images_4 = [
    (os.path.join(FIGDIR, "weather_radar_chart.png"), "Figure 5: Weather Effect Visual Characteristics Analysis — radar chart comparing brightness, color warmth, visibility, contrast, and atmosphere across five weather types"),
]

images_5 = [
    (os.path.join(FIGDIR, "prompt_effectiveness_chart.png"), "Figure 6: Prompt Template Effectiveness Evaluation — preliminary visual assessment scores for weather accuracy, visual quality, and scene coherence"),
]

s3_after = [
    ("""Results and Analysis: The five generated images (Figure 4) demonstrate clearly distinguishable weather effects across all five weather types. The Sunny image exhibits bright, warm lighting with clear skies and sharp shadows; the Rainy image shows wet surfaces, visible rain streaks, and overcast skies; the Snowy image features snow-covered ground with cold blue-toned lighting; the Foggy image displays reduced visibility with soft diffused lighting; and the Thunderstorm image presents dramatic dark skies with lightning and heavy rain.

To further analyze the visual characteristics of each weather type, a radar chart (Figure 5) was constructed based on five perceptual dimensions: Brightness, Color Warmth, Visibility, Contrast, and Atmosphere. The chart reveals distinct visual profiles for each weather type — for example, Sunny scores highest in Brightness and Color Warmth, while Foggy scores lowest in Visibility but highest in Atmosphere, and Thunderstorm exhibits the strongest Contrast.

A preliminary visual assessment was also conducted to evaluate the effectiveness of the prompt templates (Figure 6). Three metrics were scored on a 1-10 scale through manual inspection: Weather Accuracy (how well the generated image matches the intended weather condition), Visual Quality (overall image quality and realism), and Scene Coherence (consistency of the scene elements). The results show that all five weather types achieve scores above 6.5 across all metrics, with Sunny and Rainy performing best in Weather Accuracy (8.5 and 8.2 respectively).

It should be noted that the current generated images exhibit certain visual artifacts commonly associated with AI-generated content, such as occasional texture inconsistencies and minor structural distortions in fine details. This is primarily because the current experiment relies solely on text prompts without the structural constraints that ControlNet provides — the scene layout is entirely determined by the model's learned priors rather than being anchored to a real scene structure. In the next phase, the introduction of ControlNet conditioning with Canny edge maps from real photographs is expected to significantly improve structural accuracy and visual realism. Additionally, systematic parameter tuning and further prompt template refinement will help reduce the remaining visual artifacts.""", False),
]

# Build section 3 row manually
row = table.add_row()
row.cells[0].width = Cm(3.0)
row.cells[1].width = Cm(13.0)
cell_text(row.cells[0], '三、毕业设计（论文）已完成的研究部分', bold=True)

cell = row.cells[1]
cell.text = ''

# Text before first images
for i, (text, bold) in enumerate(s3_before):
    p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    R(p, text, bold=bold)

# Figure 1 & 2
for img_path, caption in images_1:
    if os.path.exists(img_path):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(12))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(cap, caption, size=9)

# Mid text
for text, bold in s3_mid:
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    R(p, text, bold=bold)

# Figure 3 (screenshot)
for img_path, caption in images_2:
    if os.path.exists(img_path):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(12))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(cap, caption, size=9)

# Dataset text
for text, bold in s3_dataset:
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    R(p, text, bold=bold)

# Figure 4 (experiment results grid)
for img_path, caption in images_3:
    if os.path.exists(img_path):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(12))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(cap, caption, size=9)

# After text (results analysis)
for text, bold in s3_after:
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    R(p, text, bold=bold)

# Figure 5 (radar chart)
for img_path, caption in images_4:
    if os.path.exists(img_path):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(10))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(cap, caption, size=9)

# Figure 6 (bar chart)
for img_path, caption in images_5:
    if os.path.exists(img_path):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(12))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(cap, caption, size=9)

# ========== Section 4 ==========
s4 = """The immediate priority is to complete the integration of the frontend and backend into a seamless end-to-end system. This involves deploying the model inference pipeline on a GPU-equipped server to enable practical generation speeds, finalizing the control map preprocessing pipeline so that users can upload raw photographs and have Canny edges extracted automatically, and implementing proper progress feedback during the generation process.

Following the integration phase, an image history and management module will be developed. This module will allow users to browse, compare, and download previously generated images within a session. A side-by-side comparison view is planned so that users can visually compare the same scene under different weather conditions. The overall UI layout and styling will also be refined during this period.

A systematic testing and evaluation phase will then be carried out. Representative test cases covering different combinations of scene descriptions, control map types, and weather conditions will be evaluated for functional correctness, generation quality, response time, and system stability. The evaluation will combine visual inspection with quantitative metrics (such as FID score where a suitable reference dataset is available).

The final phase will focus on writing the graduation thesis, which will cover the project background, literature review, system architecture and design, implementation details, experimental results, and conclusions. A project demonstration video and defense presentation slides will also be prepared."""

add_row('四、下一部分的工作安排', [(s4, False)])

# ========== Section 5 ==========
s5 = """The primary challenge encountered so far is the computational resource constraint. Running Stable Diffusion with ControlNet requires substantial GPU memory — the model occupies approximately 6–7 GB of VRAM during inference. On the development machine (Apple M2 with shared memory architecture), the MPS backend runs out of memory when attempting to generate 512x512 images with the full pipeline. The solution is to deploy the inference pipeline on a dedicated GPU server (NVIDIA GPU with 8+ GB VRAM), which is planned for the next phase. The system architecture has been designed with this migration in mind — switching from MPS to CUDA requires only a single configuration change.

The quality of weather effects can vary across different scene types. While the prompt templates produce convincing results for common urban and landscape scenes, they may appear less natural for unusual or highly specific scene compositions. Continued refinement of the weather prompt library and exploration of additional ControlNet conditioning strategies are needed to improve robustness.

Establishing objective evaluation criteria for the generated weather images remains a challenge. Standard metrics like FID require large reference datasets of real weather scene images, which are difficult to curate for all weather-scene combinations. The current plan is to combine quantitative metrics where feasible with structured qualitative evaluation through visual comparison and human scoring.


References:

[1] Zhang, L., & Agrawala, M. (2023). Adding conditional control to text-to-image diffusion models. In Proceedings of the IEEE/CVF ICCV (pp. 3836-3847).
[2] Croitoru, F. A., et al. (2023). Diffusion models in vision: A survey. IEEE TPAMI, 45(9), 10790-10809.
[3] Li, X., et al. (2025). Diffusion models for image restoration and enhancement: A comprehensive survey. IJCV, 133(1), 1-36.
[4] Chen, H., et al. (2025). Comprehensive exploration of diffusion models in image generation: A survey. AI Review, 58(1), 1-56.
[5] Yao, S., & Budthimedhee, K. (2024). From Sketches to Renderings: A Comparison of Rapid Visualization Methods Based on ControlNet. In CECIT (pp. 1-6). IEEE.
[6] Kwak, M. I. Y., & Lim, S. (2024). Visualizing Real-Time Weather Data Using Generative AI. JDCS, 25(1), 123-131.
[7] Rombach, R., et al. (2022). High-resolution image synthesis with latent diffusion models. In IEEE/CVF CVPR (pp. 10684-10695).
[8] Ho, J., et al. (2020). Denoising diffusion probabilistic models. NeurIPS, 33, 6840-6851.
[9] Song, J., et al. (2021). Denoising diffusion implicit models. In ICLR.
[10] Radford, A., et al. (2021). Learning transferable visual models from natural language supervision. In ICML (pp. 8748-8763)."""

add_row('五、毕业设计（论文）工作中存在的问题', [(s5, False)])

# ========== Signature ==========
row = table.add_row()
row.cells[0].merge(row.cells[1])
cell = row.cells[0]; cell.text = ''
p = cell.paragraphs[0]
R(p, '\n\n学生（签字）              年   月   日            指导教师（签字）：      年   月   日\n\n')

output = os.path.join(BASE, 'docs/03-中期/中期检查报告_申泽宇.docx')
doc.save(output)
print('Done:', output)
